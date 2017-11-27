#!/usr/bin/env python
# -*- coding: utf-8 -*-

from django.shortcuts import render
from django.shortcuts import redirect
from django.http import HttpResponseRedirect
from django.http import HttpResponse
from django.http import JsonResponse
from django.conf import settings
from django.contrib.auth.decorators import permission_required
from .models import Tourney, Sport, Location
from .forms import UploadForm, TourneyFormSet, TourneyForm
from .forms import LocForm, FilterForm, ImportFormSet
from .reports import export_gov, export_vfd, export_min
from .utils import week_start_date
from datetime import datetime, date, timedelta
from django.utils import formats
from django.utils.formats import localize
from django.db.models import Q, Sum
from django.forms import formset_factory
from docx import Document
import calendar
import json


def list_view(request, year=None, month=None, week=None, sport=None, only_start=True, typer='13'):
    weeks = []
    dnow = datetime.now()
    current_week = dnow.isocalendar()[1]
    select_week = next_week = prev_week = None
    custom_query = False
    typer = typer
    try:
        sport = Sport.objects.get(id=sport)
    except:
        pass
    # set filter start/end
    if year:
        year_int = int(year)
        date_start = datetime(year_int, 1, 1)
        date_end = datetime(year_int, 12, 31)
    if month:
        year_int = int(year)
        month_int = int(month)
        date_start = datetime(year_int, month_int, 1)
        last_month_day = calendar.monthrange(year_int, month_int)[1]
        date_end = datetime(year_int, month_int, last_month_day)
        first_week_month = datetime(year_int, month_int, 1).isocalendar()[1]
        if month_int == 1:
            first_week_month = 1
        last_week_month = datetime(year_int, month_int, last_month_day).isocalendar()[1]
        for week1 in range(first_week_month, last_week_month+1):
            weeks.append(week1)
    if week:
        select_week = int(week)
        next_week = select_week + 1
        prev_week = select_week - 1
        date_start = week_start_date(year_int, select_week)
        date_end = date_start + timedelta(days=6)
    if request.method == 'POST':
        form = FilterForm(request.POST)
        if form.is_valid():
            date_start = form.cleaned_data['date_start']
            date_end = form.cleaned_data['date_end']
	    typer = form.cleaned_data['typer']
            sport = form.cleaned_data['sport']
	    only_start = form.cleaned_data['only_start']
            custom_query = True
    else:
        init_data={'date_start': date_start, 'date_end': date_end, 'only_start': True}
        form = FilterForm(initial=init_data)
    # get tourney list by dates
    if sport:
        tsport = Q(sport=sport)
    if typer == '13':
        tfilt = (Q(typer='1') | Q(typer='3'))
    else:
        tfilt = Q(typer=typer)
    if only_start:
        tourney_list = Tourney.objects.filter(
            (Q(date_start__gte=date_start, date_start__lte=date_end))
        ).filter(tfilt)
    else:
        tourney_list = Tourney.objects.filter(
            (Q(date_start__lte=date_start,date_end__gte=date_end)) | 
            (Q(date_end__gte=date_start, date_end__lte=date_end)) | 
            (Q(date_start__gte=date_start, date_start__lte=date_end))
        ).filter(tfilt)
    if sport:
        tourney_list = tourney_list.filter(sport=sport)

    judje_sum = tourney_list.aggregate(
         Sum('judje_sum'),
         Sum('reward_sum'),
         Sum('print_sum'),
         Sum('proezd_sum'),
         Sum('projiv_sum'),
         Sum('pitanie_sum'),
         Sum('vznos_sum'),
         Sum('arenda_sum'),
         Sum('medic_sum'),
         Sum('transport_sum')
         )
    # export block
    if request.GET.get('file', '') == 'gov':
        return export_gov(tourney_list, date_start, date_end)
    if request.GET.get('file', '') == 'vfd':
        return export_vfd(tourney_list, date_start, date_end)
    if request.GET.get('file', '') == 'min':
        return export_min(tourney_list, date_start, date_end)
    # /end export block
    context = {'tourney_list': tourney_list,
               'date_start': date_start,
               'date_end': date_end,
               'sport': sport,
               'year': year, 'month': month, 'week': week,
               'weeks': weeks,
               'current_week': current_week,
               'next_week': next_week,
               'prev_week': prev_week,
               'form': form,
               'custom_query': custom_query,
               'judje_sum': judje_sum
               } 
    return render(request, 'sportapp/list.html', context)
    
@permission_required('sportapp.change_tourney')
def list_edit(request, year=None, month=None, week=None, sport=None, typer=1):
    if not request.user.is_authenticated:
        return redirect('%s?next=%s' % (settings.LOGIN_URL, request.path))
    sports = Sport.objects.all()
    sports_list = []
    
    if year:
        tourney_list = Tourney.objects.filter(date_start__year=year)
    if month:
        tourney_list = tourney_list.filter(date_start__month=month)
    if week:
        tourney_list = tourney_list.filter(date_start__week=week)
    if sport:
        tourney_list = tourney_list.filter(sport=sport)
    tourney_list = tourney_list.order_by("date_start")

    for sp in sports:
          value = '%s' % (sp.title)
          sp_dict = {'id': sp.id, 'label': value, 'value': value}
          sports_list.append(sp_dict)
    sports_js = json.dumps(sports_list)

    formset = TourneyFormSet(request.POST, queryset=tourney_list)
    if request.method == 'POST':
        formset = TourneyFormSet(request.POST, request.FILES, queryset=tourney_list)
        if formset.is_valid():
            formset.save()
            return redirect('list_month', year=year, month=month)
    else:
        formset = TourneyFormSet(queryset=tourney_list)
    context = {'formset': formset, 'sports_js': sports_js,
          'year': year, 'month': month, 'week': week}
    return render(request, 'sportapp/list_edit.html', context)
    
@permission_required('sportapp.change_tourney')
def edit(request,  tourney_id):
    tourney = Tourney.objects.get(id=tourney_id)
    form = TourneyForm(request.POST, instance=tourney)
    if request.method == 'POST':
        formset = TourneyForm(request.POST, instance=tourney)
        if form.is_valid():
            form.save()
            return redirect('detail', tourney_id=tourney_id)
    else:
        form = TourneyForm(instance=tourney)
    return render(request, 'sportapp/edit.html', {'form': form})

@permission_required('sportapp.delete_tourney')
def delete(request,  tourney_id):
    tourney = Tourney.objects.get(id=tourney_id)
    if request.method=='POST':
        tourney.delete()
        if 'next' in request.GET:
            return redirect(request.GET['next'])
        return redirect('list_index')
    return render(request, 'sportapp/delete.html', {'object': tourney})

def sport_loc(request):
    location = Location.objects.all()
    context = {'location':location}
    return render(request, 'sportapp/sport_loc.html', context)

def sport_loc_detail(request, id):
    location = Location.objects.get(id=id)
    tourney_loc = Tourney.objects.filter(location=location)
    context = {'location':location, 'tourney':tourney_loc}
    return render(request, 'sportapp/sport_loc_detail.html', context)

def sport_list(request):
    sports = Sport.objects.all()
    context = {'sports':sports}
    return render(request, 'sportapp/sport_list.html', context)   

def sport_detail(request, id):
    sport = Sport.objects.get(id=id) 
    tourney_list = Tourney.objects.filter(sport=sport)
    context = {'sport': sport, 'tourney_list':tourney_list}
    return render(request, 'sportapp/sport_detail.html', context)

def detail(request, tourney_id):
    tourney = Tourney.objects.get(id=tourney_id)
    context = {'tourney': tourney}
    return render(request, 'sportapp/detail.html', context)
	
def month_gov(request):
    tourney_list = Tourney.objects.order_by("date_start")
    context = {'tourney_list': tourney_list}
    return render(request, 'sportapp/month_gov.html', context)

def sport_loc_detail_edit(request,id):
    loc = Location.objects.get(id=id)
    if request.method == 'POST': 
         form = LocForm(request.POST,instance=loc)
         if form.is_valid():
           form.save()
           return HttpResponseRedirect('/')
    else:
       form = LocForm(instance=loc)
    return render(request, 'sportapp/loc_edit.html', {'form':form})

def filter(request):
    tourney_list = []
    tourney_org = []
    if request.method == 'POST':
        form = FilterForm(request.POST)
        if form.is_valid():
            date_start = form.cleaned_data['date_start']
            date_end = form.cleaned_data['date_end']
	    typer = form.cleaned_data['typer']
            tourney_list = Tourney.objects.filter(
                (Q(date_start__lte=date_start,date_end__gte=date_end)) | 
                (Q(date_end__gte=date_start, date_end__lte=date_end)) | 
                (Q(date_start__gte=date_start, date_start__lte=date_end))
            ).filter(typer=typer)
    else:
        form = FilterForm()
    return render(request, 'sportapp/filter.html', {'form': form, 'tourney_list': tourney_list})

@permission_required('sportapp.change_tourney')
def upload(request):
   # if this is a POST request we need to process the form data
    file_name = None
    document = None
    if request.method == 'POST':
        # create a form instance and populate it with data from the request:
        form = UploadForm(request.POST, request.FILES)
        # check whether it's valid:
        if form.is_valid():
            # process the data in form.cleaned_data as required
            file_name = request.FILES['file_name']
            document = Document(file_name)
            dt = datetime.now()
            for row in document.tables[2].rows:
                tnew = Tourney()
                tnew.location_id = 1
                tnew.sport_id = 1
                tnew.date_start = dt
                tnew.title = row.cells[1].text
                tnew.time_text = row.cells[2].text
                tnew.save()
            # redirect to a new URL:
            #return HttpResponseRedirect('/sportpro')

    # if a GET (or any other method) we'll create a blank form
    else:
        form = UploadForm()

    context = {'form': form, 'file': file_name, 'doc': document}
    return render(request, 'sportapp/form_upload.html', context)

@permission_required('sportapp.change_tourney')
def upload2018(request):
   # if this is a POST request we need to process the form data
    file_name = None
    document = None
    sport = None
    location = None
    tarr = []
    dm = {u'января': 1, u'февраля': 2, u'марта': 3, u'апреля': 4,
          u'мая': 5, u'июня': 6, u'июля': 7, u'августа': 8, u'сентября': 9,
          u'октября': 10, u'ноября': 11, u'декабря': 12}
    am = [u'января',u'февраля',u'марта',u'апреля',u'мая',u'июня',u'июля',u'августа',u'сентября',
          u'октября', u'ноября',u'декабря']
    if request.method == 'POST':
        # create a form instance and populate it with data from the request:
        form = UploadForm(request.POST, request.FILES)
        # check whether it's valid:
        if form.is_valid() and 'read' in request.POST:
            # process the data in form.cleaned_data as required
            file_name = request.FILES['file_name']
            sport = form.cleaned_data['sport']
            location = form.cleaned_data['location']
            document = Document(file_name)
            dt = datetime.now()
            tarr = []
            for table in document.tables:
                rowcount = 0
                for row in table.rows:
                    rowcount += 1
                    td = row.cells[2].text.split(" ")
                    if rowcount > 3 and len(row.cells[1].text)>10:
                        tnew = Tourney()
                        if form.cleaned_data['location']:
                            tnew.location = form.cleaned_data['location']
                        else: 
                            tnew.location_id = 1
                        tnew.sport = form.cleaned_data['sport']
                        #tneddw.date_start = dt
                        tnew.title = row.cells[1].text
                        #tnew.date_start = date(2018, dm[td[1]], int(td[0]))
                        tnew.date_start = date(2018, 1, 1)
                        #tnew.save()
                        tarr.append(tnew)
    # if a GET (or any other method) we'll create a blank form
    else:
        form = UploadForm()

    context = {'form': form, 'file': file_name, 'doc': document, 'tarr': tarr, 'sport': sport, 'location': location}
    return render(request, 'sportapp/form_upload.html', context)


@permission_required('sportapp.change_tourney')
def import2018(request):
    data = {'tooo':[{'title':'test1', 'date_start':'13.10.2017'},{'title':'test2', 'date_start':'15.10.2017'}]}
    if request.method == 'POST':
        idata=request.POST.get('sdata')
        data = json.loads(idata)
        sport_id = data['sport_id']
        location_id = data['location_id']
        sport = Sport.objects.get(id=sport_id)
        count = data['count']
        location = Location.objects.get(id=location_id)
        items = data['items']
        for i in range(1,count+1):
            tt = Tourney()
            tt.title = items[str(i)]['title']
            tt.date_start = items[str(i)]['date_start']
            tt.sport = sport
            tt.location = location
            tt.save()
    else:
        pass
    return JsonResponse({'accepted': True})


