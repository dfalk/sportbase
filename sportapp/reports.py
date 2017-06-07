#!/usr/bin/env python
# -*- coding: utf-8 -*-

from django.http import HttpResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as align
from docx.shared import Pt
from cStringIO import StringIO
from datetime import datetime
from django.utils import formats
from django.utils.formats import localize
import calendar
import json

def export_gov(tourney_list,date_start,date_end):
    file_name = "sportapp/reports/02-gov.docx"
    doc1 = Document(file_name)
    str_date_n = unicode(localize(date_start)).split(" ")
    t = doc1.paragraphs[2].text = ((str_date_n[0]) + ' - ' + (unicode(localize(date_end))))
    for tourney in tourney_list:  
        row1 = doc1.tables[2].add_row()
        str_date = unicode(localize(tourney.date_start)).split(" ")
        run11 = row1.cells[0].paragraphs[0].add_run(str_date[0] + " " + str_date[1])
        run11.bold = True
        run11.italic = True
        if tourney.date_end and tourney.date_end != tourney.date_start:
            end_date = unicode(localize(tourney.date_end)).split(" ")
            p11b = row1.cells[0].add_paragraph()
            p11b.alignment = align.CENTER
            run11b = p11b.add_run(" - " + end_date[0] + " " + end_date[1])
            run11b.bold = True
            run11b.italic = True
        row1.cells[0].paragraphs[0].alignment = align.CENTER
        row1.cells[1].paragraphs[0].text = tourney.title
        run21 = row1.cells[2].paragraphs[0].add_run(tourney.time_text)
        run21.bold = True
        row1.cells[2].paragraphs[0].alignment = align.CENTER
        p22 = row1.cells[2].add_paragraph(unicode(tourney.location))
        p22.alignment = align.CENTER
        if tourney.resp_gov != None:
            run31 = row1.cells[3].paragraphs[0].add_run(unicode(tourney.resp_gov))
            run31.bold = True
            row1.cells[3].paragraphs[0].alignment = align.CENTER
            pp32 = row1.cells[3].add_paragraph(unicode(tourney.resp_gov.position))
            pp32.alignment = align.CENTER
    for row in doc1.tables[2].rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(12)
    f = StringIO()
    docx_title = "gov_" + str(99) + ".docx"
    doc1.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response
    
def export_vfd(tourney_list,date_start,date_end):
    file_name = "sportapp/reports/03-vfd.docx"
    doc1 = Document(file_name)
    str_date_n = unicode(localize(date_start)).split(" ")
    t = doc1.paragraphs[2].text = ((str_date_n[0]) + ' - ' + (unicode(localize(date_end))))
    for tourney in tourney_list:
        row1 = doc1.tables[2].add_row()
        str_date = unicode(localize(tourney.date_start)).split(" ")
        run11 = row1.cells[0].paragraphs[0].add_run(str_date[0] + " " + str_date[1])
        run11.bold = True
        run11.italic = True
        if tourney.date_end and tourney.date_end != tourney.date_start:
            end_date = unicode(localize(tourney.date_end)).split(" ")
            p11b = row1.cells[0].add_paragraph()
            p11b.alignment = align.CENTER
            run11b = p11b.add_run(" - " + end_date[0] + " " + end_date[1])
            run11b.bold = True
            run11b.italic = True
        row1.cells[0].paragraphs[0].alignment = align.CENTER
        row1.cells[1].paragraphs[0].text = tourney.title
        run21 = row1.cells[2].paragraphs[0].add_run(tourney.time_vfd)
        run21.bold = True
        row1.cells[2].paragraphs[0].alignment = align.CENTER
        p22 = row1.cells[2].add_paragraph(unicode(tourney.location))
        p22.alignment = align.CENTER
        if tourney.resp_org != None:
            run31 = row1.cells[3].paragraphs[0].add_run(unicode(tourney.resp_org))
            run31.bold = True
            row1.cells[3].paragraphs[0].alignment = align.CENTER
    for row in doc1.tables[2].rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(12)
    f = StringIO()
    docx_title = "vfd_" + str(99) + ".docx"
    doc1.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response
    
def export_min(tourney_list,date_start,date_end):
    file_name = "sportapp/reports/04-min.docx"
    doc1 = Document(file_name)
    str_date_n = unicode(localize(date_start)).split(" ")
    t = doc1.paragraphs[3].text = ((str_date_n[0]) + ' - ' + (unicode(localize(date_end))))
    for tourney in tourney_list:
        # add new row
        row1 = doc1.tables[0].add_row()
        str_date = unicode(localize(tourney.date_start)).split(" ")
        run11 = row1.cells[0].paragraphs[0].add_run(str_date[0] + " " + str_date[1])
        run11.bold = True
        run11.italic = True
        # cell 1 - dates
        if tourney.date_end and tourney.date_end != tourney.date_start:
            end_date = unicode(localize(tourney.date_end)).split(" ")
            p11b = row1.cells[0].add_paragraph()
            p11b.alignment = align.CENTER
            run11b = p11b.add_run(" - " + end_date[0] + " " + end_date[1])
            run11b.bold = True
            run11b.italic = True
        row1.cells[0].paragraphs[0].alignment = align.CENTER
        # cell 2 - title
        row1.cells[1].paragraphs[0].text = tourney.title
        # cell 3 - time_text
        run21 = row1.cells[2].paragraphs[0].add_run(tourney.time_text)
        run21.bold = True
        row1.cells[2].paragraphs[0].alignment = align.CENTER
        p22 = row1.cells[2].add_paragraph(unicode(tourney.location))
        p22.alignment = align.CENTER
        # cell 4 - opening
        run21 = row1.cells[3].paragraphs[0].add_run(unicode(tourney.dt_opening))
        # cell 5 - closing
        if tourney.closing_by_end:
            run21 = row1.cells[4].paragraphs[0].add_run(u"по окончанию соревнований")
        # cell 6
        if tourney.resp_uso != None:
            run31 = row1.cells[5].paragraphs[0].add_run(unicode(tourney.resp_zam))
            run31.bold = True
            row1.cells[5].paragraphs[0].alignment = align.CENTER
            pp32 = row1.cells[5].add_paragraph("(" + unicode(tourney.resp_uso) + ")")
            pp32.alignment = align.CENTER
    # set font and etc...
    for row in doc1.tables[0].rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(12)
    # save file
    f = StringIO()
    docx_title = "min_" + str(99) + ".docx"
    doc1.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response
